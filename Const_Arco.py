import streamlit as st
import openpyxl
from docx import Document
from docx.shared import Inches
import re
import os
from datetime import datetime
import time
import matplotlib.pyplot as plt

# Cria o diretório temporário se não existir
if not os.path.exists("temp"):
    os.makedirs("temp")

def replace_text_keep_format(paragraph, old_text, new_text):
    """Substitui texto em um parágrafo mantendo a formatação."""
    for run in paragraph.runs:
        if old_text in run.text:
            updated_text = re.sub(rf'\b{re.escape(old_text)}\b', new_text, run.text)
            run.text = updated_text

def format_value(value):
    """Formata valores numéricos e datas para o padrão desejado."""
    if isinstance(value, (int, float)):
        return '{:,.2f}'.format(value).replace(',', 'X').replace('.', ',').replace('X', '.')
    elif isinstance(value, datetime):
        return value.strftime('%d/%m/%Y')
    return str(value)

def processar_constancia_arco(uploaded_excel, uploaded_word):
    """Processa os arquivos Excel e Word para gerar a Constância de Arco."""
    # Caminhos dos arquivos temporários
    excel_path = os.path.join("temp", "temp_excel.xlsx")
    word_path = os.path.join("temp", "temp_word.docx")
    img_path = os.path.join("temp", "grafico_gerado.png")
    output_word_path = os.path.join("temp", "Documento_Atualizado.docx")

    # Inicializar a barra de progresso
    progress_bar = st.progress(0)

    try:
        # Salva os arquivos carregados no diretório temporário
        with open(excel_path, "wb") as f:
            f.write(uploaded_excel.getbuffer())
        progress_bar.progress(10)

        with open(word_path, "wb") as f:
            f.write(uploaded_word.getbuffer())
        progress_bar.progress(20)

        # Carregar planilhas do Excel com OpenPyXL
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        sheet_cliente = wb['Cliente']
        sheet_max = wb['Específicos']
        sheet_fonte = wb['Fonte']
        sheet_csr = wb['CSR']
        sheet_pka = wb['Exatidão pKa']
        sheet_geo = wb['geométricos']

        # Ler os dados das células
        data = {
            "TAG1": sheet_max['A3'].value,
            "TAG2": sheet_fonte['C52'].value,
            "TAG3": sheet_fonte['D52'].value,
            "TAG4": sheet_fonte['E52'].value,
            "TAG5": sheet_fonte['F52'].value,
            "TAG6": sheet_fonte['G52'].value,
            "TAG7": sheet_fonte['C53'].value,
            "TAG8": sheet_fonte['D53'].value,
            "TAG9": sheet_fonte['E53'].value,
            "TAG10": sheet_fonte['F53'].value,
            "TAG11": sheet_fonte['G53'].value,
            "TAG12": sheet_fonte['C54'].value,
            "TAG13": sheet_fonte['D54'].value,
            "TAG14": sheet_fonte['E54'].value,
            "TAG15": sheet_fonte['F54'].value,
            "TAG16": sheet_fonte['G54'].value,
            "TAG17": sheet_fonte['D45'].value,
            "TAG18": sheet_fonte['D46'].value,
            "TAG19": sheet_fonte['D47'].value,
            "TAG20": sheet_fonte['E45'].value,
            "TAG21": sheet_fonte['E46'].value,
            "TAG22": sheet_fonte['E47'].value,
            "TAG23": sheet_fonte['F45'].value,
            "TAG24": sheet_fonte['F46'].value,
            "TAG25": sheet_fonte['F47'].value,
            "TAG26": sheet_fonte['G45'].value,
            "TAG27": sheet_fonte['G46'].value,
            "TAG28": sheet_fonte['G47'].value,
            "TAG29": sheet_fonte['H45'].value,
            "TAG30": sheet_fonte['H46'].value,
            "TAG31": sheet_fonte['H47'].value,
            "TAG32": sheet_csr['F80'].value,
            "TAG33": sheet_csr['F81'].value,
            "TAG34": sheet_csr['F82'].value,
            "TAG35": sheet_csr['G80'].value,
            "TAG36": sheet_csr['G81'].value,
            "TAG37": sheet_csr['G82'].value,
            "TAG38": sheet_csr['H80'].value,
            "TAG39": sheet_csr['H81'].value,
            "TAG40": sheet_csr['H82'].value,
            "TAG41": sheet_max['A7'].value,
            "TAG42": sheet_max['A8'].value,
            "TAG43": sheet_max['A9'].value,
            "TAG44": sheet_max['A10'].value,
            "TAG45": sheet_max['B7'].value,
            "TAG46": sheet_max['B8'].value,
            "TAG47": sheet_max['B9'].value,
            "TAG48": sheet_max['B10'].value,
            "TAG49": sheet_max['C7'].value,
            "TAG50": sheet_max['C8'].value,
            "TAG51": sheet_max['C9'].value,
            "TAG52": sheet_max['C10'].value,
            "TAG53": sheet_max['C11'].value,
            "TAG54": sheet_max['C12'].value,
            "TAG55": sheet_max['C13'].value,
            "TAG56": sheet_pka['A2'].value,
            "TAG57": sheet_pka['B2'].value,
            "TAG58": sheet_pka['C2'].value,
            "TAG59": sheet_max['C19'].value,
            "TAG60": sheet_max['D19'].value,
            "TAG61": sheet_max['B20'].value,
            "TAG62": sheet_geo['C8'].value,
            "TAG63": sheet_geo['C9'].value,
            "TAG64": sheet_geo['D8'].value,
            "TAG65": sheet_geo['A13'].value,
            "TAG66": sheet_geo['B13'].value,
            "TAG67": sheet_geo['C13'].value,

    #separação NOME TAG
    "NOME1": sheet_cliente['B1'].value,
    "NOME2": sheet_cliente['B2'].value,
    "NOME3": sheet_cliente['B3'].value,
    "NOME4": sheet_cliente['B4'].value,
    "NOME5": sheet_cliente['B5'].value,
    "NOME6": sheet_cliente['B6'].value,
    "NOME7": sheet_cliente['B7'].value,
    "NOME8": sheet_cliente['B8'].value,
    "NOME9": sheet_cliente['B9'].value,
    "NOME10": sheet_cliente['B10'].value,
    "NOME11": sheet_cliente['B11'].value,
    "NOME12": sheet_cliente['B12'].value,
    "NOME13": sheet_cliente['B16'].value,
    "NOME14": sheet_cliente['B17'].value,
    "NOME15": sheet_geo['E1'].value,
    "NOME16": sheet_geo['C4'].value,
    "NOME17": sheet_geo['E2'].value,
    "NOME18": sheet_cliente['B13'].value,
    "NOME19": sheet_cliente['B14'].value,
    "NOME20": sheet_cliente['B15'].value,
        }

        # Formatar os valores
        for key in data:
            data[key] = format_value(data[key])
        progress_bar.progress(40)

        # Obter dados para o gráfico
        sheet_graf = wb["ChartData"]  # Substitua pelo nome correto
        x_values = []
        y_values = []

        # Percorrer todas as linhas com dados nas colunas A (X) e B (Y)
        for row in sheet_graf.iter_rows(min_row=1, max_row=sheet_graf.max_row, min_col=1, max_col=2, values_only=True):
            if row[0] is not None and row[1] is not None:  # Verifica se os valores não estão vazios
                x_values.append(row[0])
                y_values.append(row[1])

        # Criar o gráfico com Matplotlib
        plt.figure(figsize=(6, 4))
        plt.plot(x_values, y_values, marker="o", linestyle="-", color="b")
        plt.xlabel("Tempo (ms)")
        plt.ylabel("Intensidade (%)")
        plt.grid(True)

        # Salvar o gráfico como imagem
        plt.savefig(img_path)
        plt.close()
        progress_bar.progress(60)

        # Abrir o documento Word e substituir os textos
        doc = Document(word_path)

        for table in doc.tables:
            for cell in table._cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        replace_text_keep_format(paragraph, key, value)

        for paragraph in doc.paragraphs:
            for key, value in data.items():
                replace_text_keep_format(paragraph, key, value)

            # Inserir o gráfico se o marcador estiver presente
            if 'INSERIR_GRAFICO' in paragraph.text:
                paragraph.text = paragraph.text.replace('INSERIR_GRAFICO', '')
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(2.5))
                break

        progress_bar.progress(80)

        # Salvar o novo documento
        doc.save(output_word_path)

        # Disponibilizar para download
        with open(output_word_path, "rb") as f:
            st.download_button("Baixar Documento Atualizado", f, file_name=f"{data['NOME1']}.docx")
        progress_bar.progress(100)

    finally:
        # Pequena pausa para garantir que o sistema libere os arquivos
        time.sleep(1)

        # Remover arquivos temporários
        for file in [excel_path, word_path, output_word_path, img_path]:
            if os.path.exists(file):
                try:
                    os.remove(file)
                except PermissionError:
                    st.warning(f"Arquivo em uso, não pode ser excluído: {file}")
