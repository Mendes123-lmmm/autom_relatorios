import streamlit as st
import openpyxl
import xlwings as xw
from docx import Document
from docx.shared import Inches
import re
import os
from datetime import datetime
import time  # Para garantir tempo suficiente para salvar o gráfico


def replace_text_keep_format(paragraph, old_text, new_text):
    """Substitui texto em um parágrafo mantendo a formatação."""
    for run in paragraph.runs:
        if old_text in run.text:
            updated_text = re.sub(rf'\b{re.escape(old_text)}\b', new_text, run.text)
            run.text = updated_text


def format_value(value):
    """Formata valores numéricos e datas para o padrão desejado."""
    if isinstance(value, (int, float)):
        return '{:,.2f}'.format(value).replace(',', 'X').replace('.', ',').replace('X', '.')
    elif isinstance(value, datetime):
        return value.strftime('%d/%m/%Y')
    return str(value)


st.title("Gerador de Documentos Automatizado")

uploaded_excel = st.file_uploader("Carregar Arquivo Excel", type=["xlsx"])
uploaded_word = st.file_uploader("Carregar Arquivo Word", type=["docx"])

if uploaded_excel and uploaded_word:
    if st.button("Gerar Documento"):
        # Criar arquivos temporários
        excel_path = "temp_excel.xlsx"
        word_path = "temp_word.docx"
        img_path = "imagem.png"
        output_word_path = "Documento_Atualizado.docx"

        # Inicializar a barra de progresso
        progress_bar = st.progress(0)

        try:
            with open(excel_path, "wb") as f:
                f.write(uploaded_excel.getbuffer())
            progress_bar.progress(10)  # Atualiza a barra de progresso

            with open(word_path, "wb") as f:
                f.write(uploaded_word.getbuffer())
            progress_bar.progress(20)  # Atualiza a barra de progresso

            # Carregar planilhas do Excel com OpenPyXL
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            sheet_cliente = wb['Cliente']
            sheet_max = wb['Específicos']
            sheet_fonte = wb['Fonte']

            # Ler os dados das células
            data = {
                "TAG1": sheet_max['A3'].value,
                "TAG2": sheet_fonte['C52'].value,
                "NOME1": sheet_cliente['B1'].value,
                "NOME2": sheet_cliente['B2'].value,
                "NOME3": sheet_cliente['B3'].value,
            }

            # Formatar os valores
            for key in data:
                data[key] = format_value(data[key])
            progress_bar.progress(40)  # Atualiza a barra de progresso

            # Exportar o gráfico usando xlwings
            # Abrir o arquivo Excel com xlwings
            wb_xlwings = xw.Book('temp_excel.xlsx')
            sheet = wb_xlwings.sheets['Gráfico']  # Substitua pelo nome da sua aba

            # Acessar o gráfico corretamente e exportá-lo como uma imagem
            chart_objects = sheet.api.ChartObjects()
            if chart_objects.Count > 0:
                chart = chart_objects.Item(1).Chart
                chart.Export(r'D:\PycharmProjects\autom_relatorios\imagem.png')
                print("Gráfico exportado com sucesso!")
            else:
                print("Nenhum gráfico encontrado na aba.")
            progress_bar.progress(60)  # Atualiza a barra de progresso

            # Abrir o documento Word e substituir os textos
            doc = Document(word_path)

            for table in doc.tables:
                for cell in table._cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            replace_text_keep_format(paragraph, key, value)

            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    replace_text_keep_format(paragraph, key, value)

                # Inserir o gráfico somente se ele foi exportado corretamente
                if 'INSERIR_GRAFICO' in paragraph.text:
                    # Substituir o termo "INSERIR_GRAFICO" pelo gráfico
                    paragraph.text = paragraph.text.replace('INSERIR_GRAFICO', '')  # Remove o marcador

                    # Inserir a imagem logo após remover o termo
                    run = paragraph.add_run()  # Cria um 'run' no mesmo parágrafo
                    run.add_picture(r'D:\PycharmProjects\autom_relatorios\imagem.png',
                                    width=Inches(2.5))  # Insere a imagem
                    break  # Interrompe o loop após inserir a imagem
            progress_bar.progress(80)  # Atualiza a barra de progresso

            # Salvar o novo documento
            doc.save(output_word_path)

            # Disponibilizar para download
            with open(output_word_path, "rb") as f:
                st.download_button("Baixar Documento Atualizado", f, file_name=f"{data['NOME1']}.docx")
            progress_bar.progress(100)  # Atualiza a barra de progresso

        finally:
            # Fechar o Excel corretamente antes de excluir arquivos
            try:
                wb_xlwings.close()  # Fecha o arquivo Excel
                xw.App().quit()
                # app.quit()  # Fecha o processo do Excel
            except NameError:
                pass  # Ignora se as variáveis não existirem

            # Pequena pausa para garantir que o sistema libere o arquivo
            time.sleep(1)

            # Remover apenas os arquivos temporários, mantendo o gráfico
            for file in [excel_path, word_path, output_word_path, img_path]:
                if os.path.exists(file):
                    try:
                        os.remove(file)
                    except PermissionError:
                        print(f"Arquivo em uso, não pode ser excluído: {file}")
